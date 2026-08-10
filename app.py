import streamlit as st
import sqlite3
import pandas as pd
import plotly.graph_objects as go
from datetime import datetime, timedelta
import subprocess
import sys
import os
import requests

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass  # python-dotenv 미설치 시 .env 로드는 건너뛰고, OS 환경변수만 사용

import db_init

# 실거래가 조회용 공공데이터포털 키 - db_init.py의 DATA_KEY와 동일한 값(같은 서비스 키)을 그대로 재사용
RTMS_KEY = db_init.DATA_KEY

# ===== 페이지 설정 =====
st.set_page_config(
    page_title="PF 사업성 검토 에이전트",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ===== 스타일 =====
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@300;400;500;700&display=swap');
    html, body, [class*="css"] { font-family: 'Noto Sans KR', sans-serif; }
    .stApp { background-color: #0f1117; }
    section[data-testid="stSidebar"] { background-color: #0f1117; }
    .header-title { font-size: 1.6rem; font-weight: 700; color: #ffffff; letter-spacing: -0.5px; }
    .header-sub { font-size: 0.85rem; color: #8b8fa8; margin-top: 2px; }
    .metric-card {
        background: #1a1d2e; border: 1px solid #2a2d3e;
        border-radius: 12px; padding: 16px 20px;
    }
    .metric-label { font-size: 0.75rem; color: #8b8fa8; margin-bottom: 4px; }
    .metric-value { font-size: 1.4rem; font-weight: 700; color: #ffffff; }
    .metric-delta { font-size: 0.78rem; margin-top: 4px; }
    .delta-up { color: #ef4444; }
    .delta-down { color: #22c55e; }
    .section-title {
        font-size: 1rem; font-weight: 600; color: #e2e8f0;
        margin: 24px 0 12px 0; border-left: 3px solid #6366f1;
        padding-left: 10px;
    }
    .filter-label { font-size: 0.78rem; color: #8b8fa8; margin-bottom: 6px; }
    .agent-box {
        background: #1a1d2e; border: 1px solid #6366f1;
        border-radius: 16px; padding: 28px 32px; margin-top: 8px;
    }
    .agent-title { font-size: 1.2rem; font-weight: 700; color: #a5b4fc; margin-bottom: 8px; }
    .agent-desc { font-size: 0.85rem; color: #8b8fa8; margin-bottom: 20px; }

    /* AI 분석 리포트(st.markdown) 텍스트 가독성 보정 */
    div[data-testid="stMarkdownContainer"] p,
    div[data-testid="stMarkdownContainer"] li,
    div[data-testid="stMarkdownContainer"] span {
        color: #e2e8f0 !important;
        line-height: 1.7;
    }
    div[data-testid="stMarkdownContainer"] h1,
    div[data-testid="stMarkdownContainer"] h2,
    div[data-testid="stMarkdownContainer"] h3 {
        color: #ffffff !important;
    }
    div[data-testid="stMarkdownContainer"] strong {
        color: #a5b4fc !important;
    }

    /* 지표 업데이트 상태 테이블 */
    .status-table-header {
        display: grid; grid-template-columns: 1.1fr 1fr 1fr 0.9fr 1.6fr 0.9fr;
        gap: 8px; padding: 8px 14px; font-size: 0.72rem; color: #8b8fa8;
        border-bottom: 1px solid #2a2d3e; letter-spacing: 0.3px;
    }
    .status-row {
        background: #1a1d2e; border: 1px solid #2a2d3e; border-radius: 10px;
        padding: 10px 14px; margin-bottom: 6px;
    }
    .status-name { font-weight: 700; color: #e2e8f0; font-size: 0.88rem; }
    .status-val { color: #cbd5e1; font-size: 0.85rem; }
    .badge {
        display: inline-block; padding: 2px 9px; border-radius: 999px;
        font-size: 0.72rem; font-weight: 700;
    }
    .badge-ok { background: rgba(34,197,94,0.15); color: #22c55e; }
    .badge-warn { background: rgba(245,158,11,0.15); color: #f59e0b; }
    .badge-na { background: rgba(139,143,168,0.15); color: #8b8fa8; }
    .status-note { color: #8b8fa8; font-size: 0.78rem; }
</style>
""", unsafe_allow_html=True)

# ===== 상수 =====
# 미분양/인구 DB의 sido 컬럼은 단축명 (서울, 경기 등)
# 인허가/착공/준공 DB의 sido 컬럼은 풀네임 (서울특별시, 경기도 등)
SIDO_FULL = {
    "서울": "서울특별시", "부산": "부산광역시", "대구": "대구광역시",
    "인천": "인천광역시", "광주": "광주광역시", "대전": "대전광역시",
    "울산": "울산광역시", "세종": "세종특별자치시", "경기": "경기도",
    "강원": "강원특별자치도", "충북": "충청북도", "충남": "충청남도",
    "전북": "전북특별자치도", "전남": "전라남도", "경북": "경상북도",
    "경남": "경상남도", "제주": "제주특별자치도"
}

DB_PATH = "pf_data.db"

def get_db():
    return sqlite3.connect(DB_PATH)

# ===== 색상 → rgba 변환 =====
COLOR_MAP = {
    "#ef4444": "rgba(239,68,68,0.1)",
    "#6366f1": "rgba(99,102,241,0.1)",
    "#f59e0b": "rgba(245,158,11,0.1)",
    "#22c55e": "rgba(34,197,94,0.1)",
    "#06b6d4": "rgba(6,182,212,0.1)",
    "#0ea5e9": "rgba(14,165,233,0.1)",
    "#a78bfa": "rgba(167,139,250,0.1)",
}

# ===== 데이터 로드 =====
@st.cache_data(ttl=300)
def load_sigungu_list():
    conn = get_db()
    df = pd.read_sql("SELECT DISTINCT sido, sigungu FROM 미분양 ORDER BY sido, sigungu", conn)
    conn.close()
    return df

@st.cache_data(ttl=300)
def load_미분양(sigungu):
    conn = get_db()
    df = pd.read_sql(
        "SELECT date, SUM(value) as value FROM 미분양 WHERE sigungu=? GROUP BY date ORDER BY date",
        conn, params=(sigungu,)
    )
    conn.close()
    return df

@st.cache_data(ttl=300)
def load_sido_data(table, sido_full):
    conn = get_db()
    df = pd.read_sql(
        f"SELECT date, value FROM [{table}] WHERE sido=? ORDER BY date",
        conn, params=(sido_full,)
    )
    conn.close()
    return df

@st.cache_data(ttl=300)
def load_인구(sigungu):
    conn = get_db()
    df = pd.read_sql(
        "SELECT date, pop, household FROM 인구 WHERE sigungu=? ORDER BY date",
        conn, params=(sigungu,)
    )
    conn.close()
    return df

@st.cache_data(ttl=300)
def load_청약(sido_short):
    conn = get_db()
    df = pd.read_sql(
        "SELECT date, supply_rate, special_rate, supply_cnt, req_cnt FROM 청약경쟁률 WHERE region=? ORDER BY date",
        conn, params=(sido_short,)
    )
    conn.close()
    return df

@st.cache_data(ttl=300)
def load_노후도(sigungu):
    conn = get_db()
    df = pd.read_sql("SELECT * FROM 노후도 WHERE sigungu=?", conn, params=(sigungu,))
    conn.close()
    return df

# ===== 실거래가 조회용 헬퍼 (온디맨드, DB 미저장) =====
@st.cache_data(ttl=3600)
def get_lawd_codes(sido_full_name, sigungu_name):
    """법정동코드 5자리(시군구코드) 목록 조회 - 실거래가 API의 LAWD_CD 파라미터용.
    반환값은 항상 리스트. 전주시/청주시/천안시처럼 구가 있는 도시는 실거래 신고 자체가
    구 단위(완산구/덕진구 등)로만 들어오므로, 통합 시 코드가 폐지 처리 안 돼 있어도
    구 코드를 항상 우선 사용한다."""
    try:
        import PublicDataReader as pdr
    except ImportError:
        return []
    code_bdong = pdr.code_bdong()

    # 1) 전주시처럼 "{sigungu_name} 완산구" 식으로 쪼개진 하위 구가 있으면 그것부터 우선 사용
    #    (상위 통합 코드는 실거래 신고 자체가 안 들어오는 레거시 코드인 경우가 많음)
    children = code_bdong[
        (code_bdong["읍면동명"] == "") &
        (code_bdong["시도명"] == sido_full_name) &
        (code_bdong["시군구명"].str.startswith(sigungu_name + " ", na=False)) &
        (code_bdong["말소일자"] == "")
    ]
    if not children.empty:
        return [str(c).strip() for c in children["시군구코드"].tolist()]

    # 2) 정확히 일치하는 활성 코드 (구가 없는 일반 시/군/구)
    exact = code_bdong[
        (code_bdong["읍면동명"] == "") &
        (code_bdong["시군구명"] == sigungu_name) &
        (code_bdong["시도명"] == sido_full_name) &
        (code_bdong["말소일자"] == "")
    ]
    if not exact.empty:
        return [str(c).strip() for c in exact["시군구코드"].tolist()]

    # 3) 세종처럼 시군구명이 빈 값인 단일 시
    sejong_like = code_bdong[
        (code_bdong["읍면동명"] == "") &
        (code_bdong["시도명"] == sido_full_name) &
        (code_bdong["시군구명"] == "") &
        (code_bdong["말소일자"] == "")
    ]
    if not sejong_like.empty:
        return [str(c).strip() for c in sejong_like["시군구코드"].tolist()]

    return []

def fetch_apt_trade_1yr(lawd_cds):
    """국토교통부 아파트매매 실거래가(상세) API - 최근 12개월 실시간 조회.
    lawd_cds는 코드 리스트 (구가 여러 개로 쪼개진 도시는 다 합쳐서 조회)
    반환값: (거래내역 DataFrame, 호출별 디버그 로그 리스트)"""
    import xmltodict
    url = "https://apis.data.go.kr/1613000/RTMSDataSvcAptTradeDev/getRTMSDataSvcAptTradeDev"

    now = datetime.now()
    cur = datetime(now.year, now.month, 1) - timedelta(days=32)  # 신고지연 감안, 전월부터 역산
    months = []
    for _ in range(12):
        months.append(cur.strftime("%Y%m"))
        cur = datetime(cur.year, cur.month, 1) - timedelta(days=2)
        cur = datetime(cur.year, cur.month, 1)
    months = sorted(set(months))

    rows = []
    debug_log = []
    for lawd_cd in lawd_cds:
        for ym in months:
            params = {
                "serviceKey": RTMS_KEY, "LAWD_CD": lawd_cd, "DEAL_YMD": ym,
                "numOfRows": "1000", "pageNo": "1"
            }
            log_entry = {"코드": lawd_cd, "월": ym, "결과": "", "건수": 0}
            try:
                res = requests.get(url, params=params, timeout=15)
                data = xmltodict.parse(res.text)
                header = data.get("response", {}).get("header", {})
                log_entry["결과"] = f"{header.get('resultCode','?')} {header.get('resultMsg','')}"
                body = data.get("response", {}).get("body", {})
                items = body.get("items")
                if not items:
                    debug_log.append(log_entry)
                    continue
                item_list = items.get("item", [])
                if isinstance(item_list, dict):
                    item_list = [item_list]
                log_entry["건수"] = len(item_list)
                for it in item_list:
                    try:
                        amt_str = str(it.get("dealAmount", "0")).replace(",", "").strip()
                        amount = float(amt_str)
                        area = float(it.get("excluUseAr", 0) or 0)
                        if area <= 0 or amount <= 0:
                            continue
                        py_price = amount / (area / 3.3058)  # 만원/평
                        rows.append({
                            "계약년월": ym,
                            "아파트": it.get("aptNm", ""),
                            "전용면적": round(area, 1),
                            "거래금액(만원)": amount,
                            "평당가": round(py_price, 1),
                            "층": it.get("floor", ""),
                            "준공년도": it.get("buildYear", "")
                        })
                    except Exception:
                        continue
            except Exception as e:
                log_entry["결과"] = f"예외: {e}"
            debug_log.append(log_entry)

    return pd.DataFrame(rows), debug_log

# ===== DB sido 컬럼 실제값 확인 =====
@st.cache_data(ttl=300)
def get_sido_samples():
    conn = get_db()
    df = pd.read_sql("SELECT DISTINCT sido FROM 인허가 LIMIT 5", conn)
    conn.close()
    return df["sido"].tolist()

# ===== 차트 함수 =====
def make_line_chart(df, x_col, y_col, title, color="#6366f1", unit="호", show_table=True):
    if df.empty:
        st.info(f"{title} 데이터가 없습니다.")
        return
    fill_color = COLOR_MAP.get(color, "rgba(99,102,241,0.1)")
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=df[x_col], y=df[y_col],
        mode="lines+markers",
        line=dict(color=color, width=2.5),
        marker=dict(size=6, color=color),
        fill="tozeroy",
        fillcolor=fill_color,
        hovertemplate=f"%{{x}}<br>{title}: %{{y:,}}{unit}<extra></extra>"
    ))
    # 날짜 형식 변환: 202506 → 2025-06
    df = df.copy()
    df[x_col] = df[x_col].astype(str).apply(lambda x: f"{x[:4]}-{x[4:6]}" if len(x) == 6 else x)

    fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=0, r=0, t=40, b=0),
        title=dict(text=title, font=dict(color="#e2e8f0", size=14)),
        xaxis=dict(showgrid=False, color="#8b8fa8", tickfont=dict(size=11), type="category"),
        yaxis=dict(showgrid=True, gridcolor="#2a2d3e", color="#8b8fa8",
                   tickfont=dict(size=11), ticksuffix=unit),
        height=320, showlegend=False,
    )
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})
    if show_table:
        table_df = df[[x_col, y_col]].copy()
        table_df.columns = ["기준월", title]
        table_df = table_df.sort_values("기준월", ascending=False).reset_index(drop=True)
        st.dataframe(table_df, use_container_width=True, hide_index=True)

def make_bar_chart(df, x_col, y_col, title, color="#a78bfa", unit="", show_table=True):
    if df.empty:
        st.info(f"{title} 데이터가 없습니다.")
        return
    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=df[x_col], y=df[y_col],
        marker_color=color,
        hovertemplate=f"%{{x}}<br>{title}: %{{y:.2f}}{unit}<extra></extra>"
    ))
    df = df.copy()
    df[x_col] = df[x_col].astype(str).apply(lambda x: f"{x[:4]}-{x[4:6]}" if len(x) == 6 else x)

    fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=0, r=0, t=40, b=0),
        title=dict(text=title, font=dict(color="#e2e8f0", size=14)),
        xaxis=dict(showgrid=False, color="#8b8fa8", tickfont=dict(size=11), type="category"),
        yaxis=dict(showgrid=True, gridcolor="#2a2d3e", color="#8b8fa8", tickfont=dict(size=11)),
        height=320, showlegend=False,
    )
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})
    if show_table:
        table_df = df[[x_col, y_col]].copy()
        table_df.columns = ["기준월", title]
        table_df = table_df.sort_values("기준월", ascending=False).reset_index(drop=True)
        st.dataframe(table_df, use_container_width=True, hide_index=True)

def latest_val(df, col="value"):
    if df.empty or col not in df.columns:
        return 0, 0
    vals = df[col].tolist()
    latest = vals[-1] if vals else 0
    prev = vals[-2] if len(vals) > 1 else latest
    return latest, latest - prev

def metric_card(label, value, delta, unit="호"):
    delta_cls = "delta-up" if delta > 0 else "delta-down"
    delta_icon = "▲" if delta > 0 else "▼"
    st.markdown(f"""
    <div class="metric-card">
        <div class="metric-label">{label}</div>
        <div class="metric-value">{int(value):,}{unit}</div>
        <div class="metric-delta {delta_cls}">{delta_icon} {abs(int(delta)):,}{unit}</div>
    </div>""", unsafe_allow_html=True)

# ===== 헤더 =====
st.markdown('<div class="header-title">🏗️ PF 사업성 검토 에이전트</div>', unsafe_allow_html=True)
st.markdown('<div class="header-sub">부동산 PF 사업성 분석을 위한 시장 지표 모니터링 시스템</div>', unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

# ===== 관리자 인증 (업데이트 버튼 권한용) =====
if "is_admin" not in st.session_state:
    st.session_state.is_admin = False

try:
    ADMIN_PASSWORD = st.secrets.get("admin_password", None)
except Exception:
    ADMIN_PASSWORD = None
if not ADMIN_PASSWORD:
    ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "changeme1234")

with st.expander("🔐 관리자 로그인" if not st.session_state.is_admin else "🔐 관리자 모드 (로그인됨)"):
    if st.session_state.is_admin:
        st.success("관리자로 로그인되어 있어요. 지표 업데이트 버튼을 사용할 수 있어요.")
        if st.button("로그아웃"):
            st.session_state.is_admin = False
            st.rerun()
    else:
        pw = st.text_input("관리자 비밀번호", type="password", key="admin_pw_input")
        if st.button("로그인"):
            if pw == ADMIN_PASSWORD:
                st.session_state.is_admin = True
                st.success("관리자 인증 완료!")
                st.rerun()
            else:
                st.error("비밀번호가 올바르지 않습니다.")

# ===== 지표별 데이터 현황 & 업데이트 =====
st.markdown('<div class="section-title">🔄 지표별 데이터 현황</div>', unsafe_allow_html=True)
if not st.session_state.is_admin:
    st.markdown(
        "<span style='color:#8b8fa8; font-size:0.82rem'>"
        "아래는 각 지표의 최신 데이터 현황이에요. 데이터 업데이트는 관리자만 가능해요.</span>",
        unsafe_allow_html=True
    )

RELEASE_NOTE = {
    "미분양": "국토부 통계누리 · 매월 25일경 공표",
    "인허가": "국토부 통계누리 · 매월 말일경 공표",
    "착공": "국토부 통계누리 · 매월 말일경 공표 (2개월 지연 반영)",
    "준공": "국토부 통계누리 · 매월 말일경 공표 (2개월 지연 반영)",
    "인구": "행안부 주민등록인구 · 매월 초순 공표",
    "청약경쟁률": "청약홈(부동산원) · 매월 25일 공표",
    "노후도": "KOSIS 인구주택총조사 · 연 1회(보통 하반기) 공표",
}
INDICATORS = ["미분양", "인허가", "착공", "준공", "인구", "청약경쟁률", "노후도"]

def fmt_ym(ym):
    if not ym:
        return "-"
    ym = str(ym)
    return f"{ym[:4]}-{ym[4:6]}" if len(ym) == 6 else ym

st.markdown(
    '<div class="status-table-header">'
    '<div>지표</div><div>DB 최신</div><div>제공처 예상 최신</div>'
    '<div>상태</div><div>다음 공개 예정</div><div>업데이트</div></div>',
    unsafe_allow_html=True
)

@st.cache_data(ttl=1800)
def cached_indicator_status(ind):
    return db_init.get_indicator_status(ind)

for ind in INDICATORS:
    status = cached_indicator_status(ind)
    db_latest = status["db_latest"]
    expected = status["expected"]
    is_current = status["is_current"]

    if ind == "노후도":
        db_disp = db_latest if db_latest else "-"
        expected_disp = "-"
        badge_html = '<span class="badge badge-na">연단위</span>'
    else:
        db_disp = fmt_ym(db_latest)
        expected_disp = fmt_ym(expected)
        if is_current:
            badge_html = '<span class="badge badge-ok">최신</span>'
        else:
            badge_html = '<span class="badge badge-warn">업데이트 필요</span>'

    c1, c2, c3, c4, c5, c6 = st.columns([1.1, 1, 1, 0.9, 1.6, 0.9])
    with c1:
        st.markdown(f'<div class="status-row"><span class="status-name">{ind}</span></div>', unsafe_allow_html=True)
    with c2:
        st.markdown(f'<div class="status-row"><span class="status-val">{db_disp}</span></div>', unsafe_allow_html=True)
    with c3:
        st.markdown(f'<div class="status-row"><span class="status-val">{expected_disp}</span></div>', unsafe_allow_html=True)
    with c4:
        st.markdown(f'<div class="status-row">{badge_html}</div>', unsafe_allow_html=True)
    with c5:
        st.markdown(f'<div class="status-row"><span class="status-note">{RELEASE_NOTE.get(ind,"")}</span></div>', unsafe_allow_html=True)
    with c6:
        btn_label = "🔄 갱신" if ind == "노후도" else "🔄 업데이트"
        btn_disabled = not st.session_state.is_admin
        clicked = st.button(
            btn_label, key=f"update_{ind}", use_container_width=True,
            disabled=btn_disabled,
            help=None if st.session_state.is_admin else "관리자만 클릭할 수 있어요. 위 '관리자 로그인'에서 인증해주세요."
        )
        if clicked:
            with st.spinner(f"{ind} 업데이트 중..."):
                child_env = os.environ.copy()
                child_env["PYTHONIOENCODING"] = "utf-8"
                result = subprocess.run(
                    [sys.executable, "db_init.py", "--update", ind],
                    capture_output=True, text=True, encoding="utf-8", errors="replace",
                    env=child_env
                )
            if result.returncode == 0:
                st.cache_data.clear()
                st.session_state["update_error"] = None
                already_current = "이미 최신 상태" in (result.stdout or "")
                st.session_state["update_success"] = {"ind": ind, "already_current": already_current}
                st.rerun()
            else:
                st.session_state["update_error"] = {
                    "ind": ind,
                    "log": (result.stderr or result.stdout or "(오류 메시지 없음)")
                }
                st.rerun()

if st.session_state.get("update_success"):
    info = st.session_state["update_success"]
    done_ind = info["ind"]
    already_current = info.get("already_current", False)

    if done_ind == "노후도":
        st.success(
            "노후도 업데이트 완료! "
            "(기준년도가 그대로라면, 아직 통계청에서 더 최신 연도 자료를 공표하지 않았다는 뜻이에요 — 정상입니다)"
        )
    elif already_current:
        st.success(f"{done_ind} 업데이트 완료! (제공처에도 이미 이 달까지만 공표되어 있어, 기존 최신월 그대로예요 — 정상입니다)")
    else:
        st.success(f"{done_ind} 업데이트 완료! 새 데이터가 추가됐어요.")
    st.session_state["update_success"] = None

if st.session_state.get("update_error"):
    err = st.session_state["update_error"]
    st.error(f"{err['ind']} 업데이트 실패")
    with st.expander("오류 로그 (전체)"):
        st.code(err["log"], language="text")

st.markdown(
    "<span style='color:#8b8fa8; font-size:0.78rem'>"
    "'업데이트'를 누르면 최고령 1개월 데이터를 지우고 최신 1개월 데이터만 추가해요 (12개월 창 유지). "
    "노후도는 통계 특성상 매년 전체를 다시 갱신해요.</span>",
    unsafe_allow_html=True
)

st.divider()

# ===== sido 실제값 확인 (디버그용 - 첫 실행시만) =====
# DB 인허가 sido 컬럼이 풀네임인지 단축명인지 자동 감지
@st.cache_data(ttl=3600)
def detect_sido_format():
    conn = get_db()
    df = pd.read_sql("SELECT DISTINCT sido FROM 인허가 LIMIT 3", conn)
    conn.close()
    if df.empty:
        return "full"
    sample = df["sido"].iloc[0]
    # 풀네임이면 "특별시", "광역시", "도" 등이 포함됨
    if any(x in sample for x in ["특별시", "광역시", "도", "자치시"]):
        return "full"
    return "short"

sido_format = detect_sido_format()

# ===== 지역 필터 =====
sigungu_df = load_sigungu_list()

st.markdown('<div class="filter-label">📍 지역 선택</div>', unsafe_allow_html=True)
col_sido, col_sgg, col_info = st.columns([1, 2, 3])

with col_sido:
    sido_list = sorted(sigungu_df["sido"].unique().tolist())
    selected_sido_short = st.selectbox("시도", sido_list, label_visibility="collapsed")

with col_sgg:
    sgg_list = sigungu_df[sigungu_df["sido"] == selected_sido_short]["sigungu"].tolist()
    selected_sigungu = st.selectbox("시군구", sgg_list, label_visibility="collapsed")

# 시도 풀네임 변환
sido_full = SIDO_FULL.get(selected_sido_short, selected_sido_short)
# DB 형식에 따라 sido 쿼리값 결정
sido_query = sido_full if sido_format == "full" else selected_sido_short

with col_info:
    st.markdown(
        f"<br><span style='color:#8b8fa8; font-size:0.85rem'>"
        f"시군구 지표: <b style='color:#e2e8f0'>{selected_sigungu}</b> &nbsp;|&nbsp; "
        f"시도 지표: <b style='color:#e2e8f0'>{selected_sido_short}</b></span>",
        unsafe_allow_html=True
    )

st.divider()

# ===== 데이터 로드 =====
df_미분양 = load_미분양(selected_sigungu)
df_인허가 = load_sido_data("인허가", sido_query)
df_착공 = load_sido_data("착공", sido_query)
df_준공 = load_sido_data("준공", sido_query)
df_인구 = load_인구(selected_sigungu)
df_청약 = load_청약(selected_sido_short)
df_노후도 = load_노후도(selected_sigungu)

# ===== 요약 카드 =====
st.markdown('<div class="section-title">📊 주요 지표 현황</div>', unsafe_allow_html=True)

m1, m2, m3, m4, m5, m6 = st.columns(6)
미분양_l, 미분양_d = latest_val(df_미분양)
인허가_l, 인허가_d = latest_val(df_인허가)
착공_l, 착공_d = latest_val(df_착공)
준공_l, 준공_d = latest_val(df_준공)
인구_l, 인구_d = latest_val(df_인구, "pop")
청약_l, 청약_d = latest_val(df_청약, "supply_rate")

with m1: metric_card(f"미분양 ({selected_sigungu})", 미분양_l, 미분양_d)
with m2: metric_card(f"인허가 ({selected_sido_short})", 인허가_l, 인허가_d)
with m3: metric_card(f"착공 ({selected_sido_short})", 착공_l, 착공_d)
with m4: metric_card(f"준공 ({selected_sido_short})", 준공_l, 준공_d)
with m5: metric_card(f"인구 ({selected_sigungu})", 인구_l, 인구_d, "명")
with m6:
    delta_cls = "delta-up" if 청약_d > 0 else "delta-down"
    delta_icon = "▲" if 청약_d > 0 else "▼"
    st.markdown(f"""
    <div class="metric-card">
        <div class="metric-label">청약경쟁률 ({selected_sido_short})</div>
        <div class="metric-value">{청약_l:.2f}:1</div>
        <div class="metric-delta {delta_cls}">{delta_icon} {abs(청약_d):.2f}</div>
    </div>""", unsafe_allow_html=True)

# ===== 지표 탭 =====
st.markdown('<div class="section-title">📈 지표별 추이</div>', unsafe_allow_html=True)

tabs = st.tabs(["📉 미분양", "🏠 인허가", "🔨 착공", "🏢 준공", "👥 인구", "🎯 청약경쟁률", "🏚️ 노후도"])

with tabs[0]:
    make_line_chart(df_미분양, "date", "value", f"미분양 추이 — {selected_sigungu}", "#ef4444", "호")

with tabs[1]:
    make_line_chart(df_인허가, "date", "value", f"인허가 추이 — {selected_sido_short}", "#6366f1", "호")

with tabs[2]:
    make_line_chart(df_착공, "date", "value", f"착공 추이 — {selected_sido_short}", "#f59e0b", "호")

with tabs[3]:
    make_line_chart(df_준공, "date", "value", f"준공 추이 — {selected_sido_short}", "#22c55e", "호")

with tabs[4]:
    col_a, col_b = st.columns(2)
    with col_a:
        make_line_chart(df_인구, "date", "pop", f"인구 추이 — {selected_sigungu}", "#06b6d4", "명", show_table=False)
    with col_b:
        make_line_chart(df_인구, "date", "household", f"세대수 추이 — {selected_sigungu}", "#0ea5e9", "세대", show_table=False)
    if not df_인구.empty:
        table_df = df_인구[["date", "pop", "household"]].copy()
        table_df.columns = ["기준월", "인구(명)", "세대수(세대)"]
        table_df = table_df.sort_values("기준월", ascending=False).reset_index(drop=True)
        st.dataframe(table_df, use_container_width=True, hide_index=True)

with tabs[5]:
    col_a, col_b = st.columns(2)
    with col_a:
        make_bar_chart(df_청약, "date", "supply_rate", f"전체 경쟁률 — {selected_sido_short}", "#a78bfa", ":1", show_table=False)
    with col_b:
        make_bar_chart(df_청약, "date", "special_rate", f"특별공급 경쟁률 — {selected_sido_short}", "#818cf8", ":1", show_table=False)
    if not df_청약.empty:
        table_df = df_청약[["date", "supply_rate", "special_rate", "supply_cnt", "req_cnt"]].copy()
        table_df.columns = ["기준월", "전체경쟁률", "특별공급경쟁률", "공급세대수", "신청건수"]
        table_df = table_df.sort_values("기준월", ascending=False).reset_index(drop=True)
        st.dataframe(table_df, use_container_width=True, hide_index=True)

with tabs[6]:
    if df_노후도.empty:
        st.info("노후도 데이터가 없습니다.")
    else:
        row = df_노후도.iloc[0]
        total = row["total"] if row["total"] > 0 else 1
        labels = ["1990년 이전", "1990~2000년", "2000~2010년", "2010년 이후"]
        values = [row["built_before_1990"], row["built_1990_2000"], row["built_2000_2010"], row["built_after_2010"]]
        colors_list = ["#ef4444", "#f59e0b", "#6366f1", "#22c55e"]

        n1, n2, n3, n4 = st.columns(4)
        for col, label, val, clr in zip([n1, n2, n3, n4], labels, values, colors_list):
            with col:
                pct = round(val / total * 100, 1)
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-label">{label}</div>
                    <div class="metric-value" style="color:{clr}">{int(val):,}호</div>
                    <div class="metric-delta" style="color:#8b8fa8">{pct}%</div>
                </div>""", unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        fig = go.Figure(go.Pie(
            labels=labels, values=values,
            marker=dict(colors=colors_list),
            hole=0.4,
            textinfo="label+percent",
            hovertemplate="%{label}<br>%{value:,}호 (%{percent})<extra></extra>"
        ))
        fig.update_layout(
            paper_bgcolor="rgba(0,0,0,0)",
            margin=dict(l=0, r=0, t=20, b=0),
            height=300,
            font=dict(color="#e2e8f0"),
            showlegend=False,
            title=dict(
                text=f"주택 노후도 현황 — {selected_sigungu} ({row['year']}년 기준)",
                font=dict(color="#e2e8f0", size=13)
            )
        )
        st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})

st.divider()

# ===== 실거래가 조회 (온디맨드, DB 미저장) =====
st.markdown('<div class="section-title">🏠 아파트 실거래가 조회 (최근 1년)</div>', unsafe_allow_html=True)
st.markdown(
    f"<span style='color:#8b8fa8; font-size:0.85rem'>"
    f"이 지표는 다른 지표와 달리 전 지역을 미리 적재하지 않고, 조회 버튼을 누른 시점에 "
    f"<b style='color:#e2e8f0'>{selected_sido_short} {selected_sigungu}</b> 지역만 실시간으로 가져와요. "
    f"국토교통부 실거래가 API를 최근 12개월치 호출해서 전용면적 평당가(거래금액 ÷ 전용면적) 추이를 계산합니다.</span>",
    unsafe_allow_html=True
)
st.markdown("<br>", unsafe_allow_html=True)

if "rtms_confirm" not in st.session_state:
    st.session_state.rtms_confirm = False
if "rtms_result" not in st.session_state:
    st.session_state.rtms_result = None
if "rtms_region" not in st.session_state:
    st.session_state.rtms_region = None

col_rtms_btn, _ = st.columns([1, 3])
with col_rtms_btn:
    if st.button("🔍 실거래가 조회", use_container_width=True, type="primary", key="rtms_ask"):
        st.session_state.rtms_confirm = True

if st.session_state.rtms_confirm:
    st.warning(
        f"**{selected_sido_short} {selected_sigungu}**의 최근 1년치 실거래가를 조회하시겠습니까? "
        f"(API를 12회 호출하며 완료까지 몇 초 정도 걸릴 수 있어요)"
    )
    c_yes, c_no = st.columns([1, 1])
    with c_yes:
        do_fetch = st.button("✅ 예, 조회할게요", use_container_width=True, key="rtms_yes")
    with c_no:
        cancel = st.button("취소", use_container_width=True, key="rtms_no")

    if cancel:
        st.session_state.rtms_confirm = False
        st.rerun()

    if do_fetch:
        lawd_cds = get_lawd_codes(sido_full, selected_sigungu)
        if not lawd_cds:
            st.error("이 지역의 법정동코드(LAWD_CD)를 찾을 수 없습니다. `PublicDataReader` 패키지가 설치되어 있는지 확인해주세요.")
        else:
            with st.spinner(f"{selected_sigungu} 최근 12개월 실거래가 수집 중..."):
                trade_df, debug_log = fetch_apt_trade_1yr(lawd_cds)
            st.session_state.rtms_result = trade_df
            st.session_state.rtms_region = f"{selected_sido_short} {selected_sigungu}"
            st.session_state.rtms_debug = {"codes": lawd_cds, "log": debug_log}
        st.session_state.rtms_confirm = False
        st.rerun()

if st.session_state.rtms_result is not None:
    trade_df = st.session_state.rtms_result
    region_label = st.session_state.rtms_region
    debug_info = st.session_state.get("rtms_debug")
    if debug_info:
        with st.expander("🔧 조회 디버그 정보 (문제 확인용)"):
            st.write("사용된 법정동코드:", debug_info.get("codes"))
            st.dataframe(pd.DataFrame(debug_info.get("log", [])), use_container_width=True, hide_index=True)
    if trade_df.empty:
        st.info(f"{region_label} 최근 1년간 조회된 아파트 매매 거래가 없습니다.")
    else:
        monthly = (
            trade_df.groupby("계약년월")["평당가"]
            .agg(평당가="mean", 거래건수="count")
            .reset_index()
            .sort_values("계약년월")
        )
        make_line_chart(
            monthly.rename(columns={"계약년월": "date", "평당가": "value"}),
            "date", "value", f"전용면적 평당가 추이 — {region_label}", "#ef4444", "만원", show_table=False
        )

        col_a, col_b = st.columns(2)
        with col_a:
            latest_price = monthly["평당가"].iloc[-1] if not monthly.empty else 0
            st.markdown(f"""
            <div class="metric-card">
                <div class="metric-label">최신월 평균 평당가</div>
                <div class="metric-value">{latest_price:,.0f}만원</div>
            </div>""", unsafe_allow_html=True)
        with col_b:
            total_cnt = int(monthly["거래건수"].sum())
            st.markdown(f"""
            <div class="metric-card">
                <div class="metric-label">최근 1년 총 거래건수</div>
                <div class="metric-value">{total_cnt:,}건</div>
            </div>""", unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        # ===== 준공년도별 전용평당가 =====
        year_df = trade_df.copy()
        year_df["준공년도"] = pd.to_numeric(year_df["준공년도"], errors="coerce")
        year_df = year_df.dropna(subset=["준공년도"])
        year_df["준공년도"] = year_df["준공년도"].astype(int)

        if not year_df.empty:
            by_year = (
                year_df.groupby("준공년도")["평당가"]
                .agg(평당가="mean", 거래건수="count")
                .reset_index()
                .sort_values("준공년도")
            )
            by_year["준공년도"] = by_year["준공년도"].astype(str)

            fig_year = go.Figure()
            fig_year.add_trace(go.Bar(
                x=by_year["준공년도"], y=by_year["평당가"],
                marker_color="#f59e0b",
                customdata=by_year["거래건수"],
                hovertemplate="%{x}년 준공<br>평균 평당가: %{y:,.0f}만원<br>거래건수: %{customdata}건<extra></extra>"
            ))
            fig_year.update_layout(
                paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                margin=dict(l=0, r=0, t=40, b=0),
                title=dict(text=f"준공년도별 전용평당가 — {region_label}", font=dict(color="#e2e8f0", size=14)),
                xaxis=dict(showgrid=False, color="#8b8fa8", tickfont=dict(size=11), type="category"),
                yaxis=dict(showgrid=True, gridcolor="#2a2d3e", color="#8b8fa8",
                           tickfont=dict(size=11), ticksuffix="만원"),
                height=320, showlegend=False,
            )
            st.plotly_chart(fig_year, use_container_width=True, config={"displayModeBar": False})

            year_table = by_year.rename(columns={"평당가": "평균 평당가(만원)", "거래건수": "거래건수"})
            year_table["평균 평당가(만원)"] = year_table["평균 평당가(만원)"].round(0)
            year_table = year_table.sort_values("준공년도", ascending=False).reset_index(drop=True)
            st.dataframe(year_table, use_container_width=True, hide_index=True)

        st.markdown("<br>", unsafe_allow_html=True)
        detail_df = trade_df[["계약년월", "아파트", "전용면적", "거래금액(만원)", "평당가", "층"]].copy()
        detail_df.columns = ["계약월", "단지명", "전용면적(㎡)", "거래금액(만원)", "평당가(만원)", "층"]
        detail_df = detail_df.sort_values("계약월", ascending=False).reset_index(drop=True)
        st.dataframe(detail_df, use_container_width=True, hide_index=True)

st.divider()

# ===== PF 매력도 분석 에이전트 =====
st.markdown('<div class="section-title">🤖 지역 PF 매력도 분석 에이전트</div>', unsafe_allow_html=True)
st.markdown("""
<div class="agent-box">
<div class="agent-title">📋 AI 기반 PF 사업성 분석 리포트 생성</div>
<div class="agent-desc">수집된 데이터를 기반으로 선택 지역의 PF 사업성 분석 리포트를 Word 문서로 생성합니다. (Gemini AI 사용)</div>
</div>
""", unsafe_allow_html=True)

rtms_ready = (
    st.session_state.get("rtms_result") is not None and
    st.session_state.get("rtms_region") == f"{selected_sido_short} {selected_sigungu}"
)

st.markdown("<br>", unsafe_allow_html=True)
col_input, col_btn = st.columns([3, 1])
with col_input:
    target_region = st.text_input(
        "분석 지역 입력",
        value=f"{selected_sido_short} {selected_sigungu}",
        placeholder="예: 경기 화성시 동탄2신도시",
        label_visibility="collapsed"
    )
with col_btn:
    run_analysis = st.button(
        "📋 분석 리포트 생성", use_container_width=True, type="primary",
        disabled=not rtms_ready
    )

if not rtms_ready:
    st.caption("⚠️ 분석 리포트는 위에서 **현재 선택된 지역**의 실거래가를 먼저 조회해야 생성할 수 있어요. (지역을 바꾸면 그 지역 실거래가를 다시 조회해야 해요)")

if run_analysis:
    with st.spinner(f"**{target_region}** PF 매력도 분석 중... (30초~1분 소요)"):
        summary = {
            "분석지역": target_region,
            "미분양(최신)": f"{int(미분양_l):,}호",
            "미분양(전월대비)": f"{int(미분양_d):+,}호",
            "인허가(최신)": f"{int(인허가_l):,}호",
            "착공(최신)": f"{int(착공_l):,}호",
            "준공(최신)": f"{int(준공_l):,}호",
            "인구(최신)": f"{int(인구_l):,}명",
            "인구(전월대비)": f"{int(인구_d):+,}명",
            "청약경쟁률(최신)": f"{청약_l:.2f}:1",
        }
        if not df_노후도.empty:
            row = df_노후도.iloc[0]
            total = row["total"] if row["total"] > 0 else 1
            summary["노후주택비율(1990년이전)"] = f"{round(row['built_before_1990']/total*100,1)}%"

        rtms_df = st.session_state.get("rtms_result")
        if rtms_df is not None and not rtms_df.empty:
            summary["실거래가(최신월 평균 평당가)"] = f"{rtms_df.groupby('계약년월')['평당가'].mean().sort_index().iloc[-1]:,.0f}만원"
            summary["실거래가(최근1년 거래건수)"] = f"{len(rtms_df):,}건"

        prompt = f"""당신은 부동산 PF 사업성 검토 전문가입니다.
아래 실제 수집 데이터를 기반으로 {target_region} 지역의 PF 사업성 매력도 분석 리포트를 작성해주세요.

[수집 데이터 요약]
{chr(10).join([f"- {k}: {v}" for k, v in summary.items()])}

[리포트 구성]
# {target_region} PF 사업성 매력도 분석 리포트

## 1. 종합 평가
(★ 5점 척도 + 종합 한줄 평가)

## 2. 수요 분석
(인구 추이, 청약경쟁률 기반 수요 강도 분석)

## 3. 공급 분석
(인허가/착공/준공/미분양 기반 공급 압력 분석)

## 4. 리스크 요인
(주요 리스크 3가지 이상)

## 5. 투자 포인트 및 권고사항
(구체적 수치 근거 포함)

각 항목을 구체적 수치와 함께 전문적으로 작성해주세요."""

        try:
            import google.generativeai as genai

            GEMINI_KEY = os.environ.get("GEMINI_API_KEY", "")
            if not GEMINI_KEY:
                st.error(".env 파일에 GEMINI_API_KEY가 설정되어 있지 않습니다.")
                st.stop()

            genai.configure(api_key=GEMINI_KEY)
            model = genai.GenerativeModel("gemini-2.5-flash")
            response = model.generate_content(prompt)
            report_text = response.text

            st.markdown("---")
            st.markdown("### 📄 분석 결과")
            st.markdown(report_text)

            # Word 보고서 생성
            try:
                from docx import Document
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                import io

                doc = Document()
                title_para = doc.add_heading("PF 사업성 매력도 분석 리포트", 0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sub_para = doc.add_paragraph(
                    f"대상 지역: {target_region}  |  작성일: {datetime.now().strftime('%Y년 %m월 %d일')}"
                )
                sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph("")

                doc.add_heading("수집 데이터 요약", 2)
                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "지표"
                hdr[1].text = "값"
                for k, v in summary.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = k
                    row_cells[1].text = str(v)
                doc.add_paragraph("")

                for line in report_text.split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith("## "):
                        doc.add_heading(line.replace("## ", ""), 2)
                    elif line.startswith("# "):
                        doc.add_heading(line.replace("# ", ""), 1)
                    elif line.startswith("- ") or line.startswith("• "):
                        doc.add_paragraph(line[2:], style="List Bullet")
                    else:
                        doc.add_paragraph(line)

                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)

                st.download_button(
                    label="📥 Word 보고서 다운로드 (.docx)",
                    data=buf,
                    file_name=f"PF분석_{target_region.replace(' ','_')}_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )
            except ImportError:
                st.warning("Word 생성을 위해 터미널에서 `pip install python-docx` 실행 후 다시 시도해주세요.")

        except ImportError:
            st.error("Gemini 패키지가 없습니다. 터미널에서 `pip install google-generativeai` 실행해주세요.")
        except Exception as e:
            st.error(f"분석 중 오류 발생: {e}")
            st.info(".env 파일에 GEMINI_API_KEY가 올바르게 설정되어 있는지 확인해주세요.")
